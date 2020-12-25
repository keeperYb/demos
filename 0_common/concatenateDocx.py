from docx import Document
from docx import document

file1 = r'C:\Users\xuyb\Desktop\word\word_concatenate\module1.docx'
file2 = r'C:\Users\xuyb\Desktop\word\word_concatenate\module2.docx'


def getAllFormat():
    """读取docx的格式"""
    docx1 = Document(file1)
    for para in docx1.paragraphs:
        print(para.style.name)
    pass


def printAllStyles():
    docx = Document()

    styles = the_docx.styles
    for style in styles:
        print(style)
    pass


def getHeader():
    """得到页眉"""
    docx1 = Document(file1)
    section = docx1.sections[0]  # Fruit? [0]?
    header = section.header
    print(header.paragraphs[0].text)
    pass


def getFooter():
    docx1 = Document(file1)
    section = docx1.sections[0]
    footer = section.footer
    print(footer.paragraphs[0].text)
    pass


def addHeadings(doc: document.Document):
    """加标题, 从lv1 到 lv9"""
    for lvl in range(1, 10):
        doc.add_heading("A LEVEL " + str(lvl) + " HEADING", level=lvl)
    pass


def generate_docx(doc):
    assert isinstance(doc, document.Document)
    doc.add_paragraph('Lorem ipsum dolor sit amet.', style='List Bullet')
    pass


def generateTable(doc: document.Document):
    """表格设置"""
    table = doc.add_table(2, 2)
    table.style = 'LightShading-Accent1'
    pass


if __name__ == '__main__':
    the_docx = Document()
    assert isinstance(the_docx, document.Document)
    generateTable(the_docx)
    the_docx.save(file1)
    pass
